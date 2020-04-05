import requests
import os
import sys
import time
from bs4 import BeautifulSoup
from selenium import webdriver
from sys import argv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from views.allnews import newslink
from individual_views.baahrakhari import barakhari


class DataMining:
	fail = 'No Connection, Please Check Your Internet Connection!'
	pl = 0

	def __init__(self,siteurl,newscount):
		try:
			self.page = requests.get(siteurl)
			self.sitehome = siteurl
			self.numberofnews = newscount
		except (requests.exceptions.Timeout, requests.exceptions.TooManyRedirects, requests.exceptions.RequestException, requests.exceptions.ConnectionError, requests.exceptions.HTTPError,):
			self.page = 'noconnection'
			print self.noConn()
			sys.exit()
		return None

	def noConn(self):
		return self.fail

	def otherPage(self,site_new_url):
		try:
			self.othp = requests.get(site_new_url)
		except( requests.exceptions.TooManyRedirects):
			time.sleep(5)
		return self.othp

	def splitUrl(self,geturl):
		surl = geturl.split('/')
		return surl.pop()

	def homeUrl(self,gethm):
		return gethm.rstrip('/')

	def checkusnicode(self,s):
		if isinstance(s, str):
			link_type = 1
		elif isinstance(s, unicode):
			link_type = 0
		return link_type

	def getportalname(self,url):
		self.splittedurl = url.split('.com/')
		return self.splittedurl[0].replace('https://','')

	def mainScrape(self):
		if(self.page != 'noconnection'):
			if(self.page.status_code == 200):
				soup = BeautifulSoup(self.page.content, 'html.parser')
				# title = self.moreNews(soup)
				if(self.sitehome == 'https://onlinekhabar.com/'):
					nav = self.navItem(soup,'div','menu-primary-menu-container')
				elif(self.sitehome == 'https://gorkhapatraonline.com/'):
					nav = self.navItem(soup,'ul','navbar-nav')	
				elif(self.sitehome == 'https://imagekhabar.com/'):
					nav = self.navItem(soup,'div','ne-main-menu')				
			else:
				return self.noConn()
		return None

	def filter(self,allcont,htm_element,clss):
		return allcont.find(htm_element,class_=clss)

	def checknotnone(self,allcont,htm_element,clss):
		if allcont.find(htm_element,class_=clss) is not None:
			return allcont.find(htm_element,class_=clss)

		return None 

	def readallnews(self,getMainCat):
		newsoupsingle = BeautifulSoup(self.otherPage(self.sitehome + 'all%s' %getMainCat).content, 'html.parser')
		self.category(newsoupsingle,getMainCat,'','')
		return None

	def preventloop(self,Val):
		global pl
		pl = Val
		return self.pl

	def paginationcall(self,soup,MainCategory,SubCategory,count):
		for page in soup:
			if count > int(self.numberofnews):
				break
			elif page.has_attr('href'):
				pagesinglesoup = BeautifulSoup(self.otherPage(page['href']).content, 'html.parser')
				self.moreNews(pagesinglesoup,MainCategory,SubCategory,count)
		return None

	def moreNews(self,morecontent,MainCategory,SubCategory,counter):
		single_page = []
		outofcurrentpage = 0;
		more_news = morecontent.find_all('a',{'class':'title__regular'})
		pagination = morecontent.find_all('a',{'class':'next page-numbers'})
		if more_news is not None:
			count = counter
			for single in more_news:
				if count > int(self.numberofnews):
					outofcurrentpage = 1
					break
				elif single.has_attr('href'):
					soupsingle = BeautifulSoup(self.otherPage(single['href']).content, 'html.parser')
					self.saveindividual(soupsingle,MainCategory,SubCategory,count)
				count = count + 1
		if outofcurrentpage == 0:
			if pagination is not None:
				self.paginationcall(pagination,MainCategory,SubCategory,count)
		return None

	def saveindividual(self,allcont,MainCategory,SubCategory,counter):
		if allcont is not None:
			document = Document()
			if(self.sitehome == 'https://onlinekhabar.com/'):
				getsoupsiglehead = self.checknotnone(allcont,'h2', 'mb-0')
				getsoupsigle = self.checknotnone(allcont,'div', 'main__read--content')
			elif(self.sitehome == 'https://gorkhapatraonline.com/'):
				getsoupsiglehead = self.checknotnone(allcont,'h1', 'post-title')
				getsoupsigle = self.checknotnone(allcont,'div', 'newstext')
			elif(self.sitehome == 'https://imagekhabar.com/'):
				getsoupsiglehead = self.checknotnone(allcont,'h1', 'title-semibold-dark size-c30')
				getsoupsigle = self.checknotnone(allcont,'div', 'news-details-layout1')
			sn = self.getportalname(self.sitehome)
			if getsoupsigle is not None:
				singlecontent = getsoupsigle.select('p')
				h = getsoupsiglehead.get_text()
				content = document.add_heading(getsoupsiglehead.get_text(),2)
				for cont in singlecontent:
					conten1 = document.add_paragraph(cont.get_text())
				path = os.getcwd()

			if not os.path.exists('News'):
				os.mkdir('News')
				if not os.path.exists('News/%s'% sn):
					os.mkdir('News/%s' %sn)
					if not os.path.exists('News/%s/%s' %(sn,MainCategory)):
						os.mkdir('News/%s/%s' %(sn,MainCategory))
						if(SubCategory != ''):
							if not os.path.exists('News/%s/%s/%s' %(sn,MainCategory, SubCategory)):
								os.mkdir('News/%s/%s' %(MainCategory, SubCategory))
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
						elif(SubCategory == ''):
							document.save('News/%s/%s/%s_news_%s.docx' %(sn,MainCategory, MainCategory, counter))
			else:
				if not os.path.exists('News/%s' % sn):
					os.mkdir('News/%s' %sn)
					if not os.path.exists('News/%s/%s' % (sn,MainCategory)):
						os.mkdir('News/%s/%s' % (sn,MainCategory))
						if (SubCategory != ''):
							if not os.path.exists('News/%s/%s/%s' %(sn,MainCategory, SubCategory)):
								os.mkdir('News/%s/%s/%s' %(sn,MainCategory, SubCategory))
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
							else:
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
						elif(SubCategory == ''):
							document.save('News/%s/%s/%s_news_%s.docx' %(sn,MainCategory, MainCategory, counter))
					else:
						if (SubCategory != ''):
							if not os.path.exists('News/%s/%s/%s' %(sn,MainCategory, SubCategory)):
								os.mkdir('News/%s/%s/%s' %(sn,MainCategory, SubCategory))
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
							else:
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
						elif(SubCategory == ''):
							document.save('News/%s/%s/%s_news_%s.docx' %(sn,MainCategory, MainCategory, counter))
				else:
					if not os.path.exists('News/%s/%s' % (sn,MainCategory)):
						os.mkdir('News/%s/%s' % (sn,MainCategory))
						if (SubCategory != ''):
							if not os.path.exists('News/%s/%s/%s' %(sn,MainCategory, SubCategory)):
								os.mkdir('News/%s/%s/%s' %(sn,MainCategory, SubCategory))
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
							else:
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
						elif(SubCategory == ''):
							document.save('News/%s/%s/%s_news_%s.docx' %(sn,MainCategory, MainCategory, counter))
					else:
						if (SubCategory != ''):
							if not os.path.exists('News/%s/%s/%s' %(sn,MainCategory, SubCategory)):
								os.mkdir('News/%s/%s/%s' %(sn,MainCategory, SubCategory))
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
							else:
								document.save('News/%s/%s/%s/%s_news_%s.docx' %(sn,MainCategory, SubCategory, SubCategory, counter))
						elif(SubCategory == ''):
							document.save('News/%s/%s/%s_news_%s.docx' %(sn,MainCategory, MainCategory, counter))

		return None

	def category(self,getsoup,MainCatName,SubCategory,counter):
		if(self.sitehome == 'https://onlinekhabar.com/'):
			if(MainCatName == 'news'):
				pass
			else:
				selectCat = getsoup.find_all('a',{'class':'read__all--dot'});
				# print MainCatName
				for individualcat in selectCat:
					if individualcat.has_attr('href'):
						newUrlcat =  str(self.homeUrl(self.sitehome) + individualcat['href'] )
						catsoup = BeautifulSoup(self.otherPage(newUrlcat).content, 'html.parser')
						self.moreNews(catsoup,MainCatName,self.splitUrl(str(individualcat['href'])),1)

		elif (self.sitehome == 'https://gorkhapatraonline.com/'):
			global pl
			if getsoup is not None:
				linklist =[]
				selectCat = getsoup.find('div',{'class':'sports-groups'})
				nextpagelk = getsoup.find('a',{'rel':'next'})
				if selectCat and nextpagelk is not None:
					# print MainCatName
					getlink = selectCat.find_all('a')
					outofcurrentpage = 0
					if counter != '':
						count = counter
					else:
						count = 1
					for indlk in getlink:
						if count > int(self.numberofnews):
							outofcurrentpage = 1
							self.preventloop(0)
							break
						else:
							if indlk.has_attr('href'):
								verifiedlk = indlk['href'].replace('www.','')
								catsoup = BeautifulSoup(self.otherPage(verifiedlk).content, 'html.parser')
								if(SubCategory != ''):
									# print SubCategory
									self.saveindividual(catsoup,MainCatName,SubCategory,count)
								else:
									self.saveindividual(catsoup,MainCatName,'',count)
						count = count + 1

					if outofcurrentpage == 0:
						newsoup = BeautifulSoup(self.otherPage(nextpagelk['href']).content, 'html.parser')
						if(SubCategory != ''):
							self.category(newsoup,MainCatName,SubCategory,count)
						else:
							self.category(newsoup,MainCatName,'',count)
					

				elif pl != 1:
					if (MainCatName == 'province'):
						for i in range(1,7):
							self.preventloop(1)
							# print self.sitehome + MainCatName+'%s' %i
							newsoupsingle = BeautifulSoup(self.otherPage(self.sitehome + 'Province' + '%s' %i).content, 'html.parser')
							self.category(newsoupsingle,MainCatName,MainCatName+'%s' %i,'')

				elif (selectCat and nextpagelk) is None and MainCatName != 'nayanepal' and pl != 1:
					self.readallnews(MainCatName)

		elif (self.sitehome == 'https://imagekhabar.com/'):
			if getsoup is not None:
				linklist =[]
				selectCat = getsoup.find('div',{'class':'col-lg-8 col-md-12'})
				nextpagelk = getsoup.find('a',{'rel':'next'})
				# print nextpagelk['href']
				if selectCat and nextpagelk is not None:
					getlink = selectCat.find_all('a',{'class':'img-opacity-hover img-overlay-70'})
					# print getlink
					outofcurrentpage = 0
					if counter != '':
						count = counter
					else:
						count = 1
					for indlk in getlink:
						if count > int(self.numberofnews):
							outofcurrentpage = 1
							break
						else:
							if indlk.has_attr('href'):
								newlink = self.homeUrl(self.sitehome) + indlk['href']
								verifiedlk = newlink.replace('www.','')
								# print newlink
								catsoup = BeautifulSoup(self.otherPage(verifiedlk).content, 'html.parser')
								self.saveindividual(catsoup,MainCatName,'',count)
						count = count + 1

					if outofcurrentpage == 0:
						# print nextpagelk['href']
						newsoup = BeautifulSoup(self.otherPage(str(nextpagelk['href'])).content, 'html.parser')
						# print newsoup
						if(SubCategory != ''):
							self.category(newsoup,MainCatName,SubCategory,count)
						else:
							self.category(newsoup,MainCatName,'',count)
					

		return None

	def navItem(self,allcont,htmele,eleclass):
		nav_strip = self.filter(allcont,htmele,eleclass)
		get_nav_item = nav_strip.select('a')
		nav_link = []
		cat = []
		for a in get_nav_item:
			if a.has_attr('href') and a['href'] != '/' and a['href'] != '#':
				cat.append(self.splitUrl(str(a['href'].encode('ascii', 'ignore').decode('ascii'))))
				nav_link.append(str(a['href'].encode('ascii', 'ignore').decode('ascii')))
			else:
				continue
		for nav,cattitle in zip(nav_link,cat):
			if(self.sitehome == 'https://onlinekhabar.com/'):
				newsoup = BeautifulSoup(self.otherPage(nav).content, 'html.parser')
				self.category(newsoup,cattitle,'','');
			elif (self.sitehome == 'https://gorkhapatraonline.com/'):
				newsoup = BeautifulSoup(self.otherPage(nav).content, 'html.parser')
				self.category(newsoup,cattitle,'','')
			elif (self.sitehome == 'https://imagekhabar.com/'):
				if(nav != 'http://archive.imagekhabar.com/'):
					newsoup = BeautifulSoup(self.otherPage(self.homeUrl(self.sitehome) + nav).content, 'html.parser')
					self.category(newsoup,cattitle,'','')
		return None


